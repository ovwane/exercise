import docx
from docx.oxml.ns import qn
from docx.shared import Pt

doc = docx.Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.add_paragraph(style="Title").add_run(u"测试平台接口")
doc.add_paragraph(style="Heading 1").add_run(u"1.用户").font.size = 300000
doc.add_paragraph(style="Heading 2").add_run("登录", style="Default Paragraph Font")
doc.add_paragraph("简要描述：", style="Body Text").paragraph_format.space_before = Pt(14)
doc.add_paragraph("用户登录接口：", style="List Paragraph")
doc.add_paragraph("请求URL：", style="Body Text")
doc.add_paragraph("http://xxxx/api/user/login", style="List Paragraph")
doc.add_paragraph("请求方式：", style="Body Text")
doc.add_paragraph("POST：", style="List Paragraph")
doc.add_paragraph("参数：", style="Body Text")
table=doc.add_table(rows=5, cols=4, style="Medium Shading 1 Accent 1")
for row in table.rows:
    row.cells[0].text = 'password'
    row.cells[1].text = 'int'
    row.cells[2].text = '必填'
    row.cells[3].text = 'ad'
doc.add_paragraph("返回示例：", style="Body Text")
doc.add_paragraph("：", style="Normal")
doc.add_paragraph("返回：", style="Body Text")
table=doc.add_table(rows=3, cols=3, style="Medium Shading 1 Accent 1")
for row in table.rows:
    row.cells[0].text = '1'
    row.cells[1].text = '1'
    row.cells[2].text = '1'
doc.save("接口测试文档.docx")
from docx import *
from docx.enum.style import WD_STYLE_TYPE

document = Document()
styles = document.styles

# 生成所有表样式
for s in styles:
    if s.type == WD_STYLE_TYPE.TABLE:
        document.add_paragraph("表格样式 :  " + s.name)
        table = document.add_table(3, 3, style=s)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '第一列内容'
        heading_cells[1].text = '第二列内容'
        heading_cells[2].text = '第三列内容'
        document.add_paragraph("\n")

document.save('demo2.docx')